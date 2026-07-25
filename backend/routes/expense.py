# routes/expense.py - COMPLETE FILE WITH SHOP FILTERING
from flask import request, jsonify, send_file
from flask_login import login_required, current_user
from extensions import db
from models.expense import Expense
from datetime import datetime
import io
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn, OxmlElement

def get_current_shop_id():
    """Get the current shop ID from the logged-in user"""
    if hasattr(current_user, 'id'):
        return current_user.id
    return None

def init_expense_routes(app):
    
    # ============ EXPENSE ROUTES ============
    
    @app.route('/api/expenses', methods=['GET'])
    @login_required
    def get_expenses():
        """Get all expenses for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            status = request.args.get('status')
            payment_method = request.args.get('payment_method')
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            search = request.args.get('search')
            
            # Start with shop filter
            query = Expense.query.filter_by(shop_id=shop_id)
            
            if status:
                query = query.filter_by(status=status)
            
            if payment_method:
                query = query.filter_by(payment_method=payment_method)
            
            if start_date:
                query = query.filter(Expense.date >= start_date)
            
            if end_date:
                query = query.filter(Expense.date <= end_date)
            
            if search:
                query = query.filter(Expense.item_name.contains(search))
            
            expenses = query.order_by(Expense.date.desc()).all()
            return jsonify([e.to_dict() for e in expenses])
            
        except Exception as e:
            print(f"Error fetching expenses: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/expenses/<int:expense_id>', methods=['GET'])
    @login_required
    def get_expense(expense_id):
        """Get a specific expense for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            expense = Expense.query.filter_by(id=expense_id, shop_id=shop_id).first()
            if not expense:
                return jsonify({'error': 'Expense not found'}), 404
            
            return jsonify(expense.to_dict())
            
        except Exception as e:
            print(f"Error fetching expense: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/expenses', methods=['POST'])
    @login_required
    def create_expense():
        """Create a new expense for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            data = request.get_json()
            
            required = ['item_name', 'amount', 'date', 'payment_method']
            missing = [f for f in required if f not in data]
            if missing:
                return jsonify({'error': f'Missing required fields: {", ".join(missing)}'}), 400
            
            try:
                amount = float(data['amount'])
                if amount < 0:
                    return jsonify({'error': 'Amount cannot be negative'}), 400
            except ValueError:
                return jsonify({'error': 'Invalid amount format'}), 400
            
            # Create expense with shop_id
            expense = Expense(
                shop_id=shop_id,
                item_name=data['item_name'],
                amount=amount,
                date=datetime.strptime(data['date'], '%Y-%m-%d').date(),
                payment_method=data['payment_method'],
                reference=data.get('reference', ''),
                status=data.get('status', 'Pending'),
                notes=data.get('notes', ''),
                created_by=current_user.id if current_user.is_authenticated else None
            )
            
            db.session.add(expense)
            db.session.commit()
            
            return jsonify(expense.to_dict()), 201
            
        except Exception as e:
            db.session.rollback()
            print(f"Error creating expense: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/expenses/<int:expense_id>', methods=['PUT'])
    @login_required
    def update_expense(expense_id):
        """Update an expense for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            expense = Expense.query.filter_by(id=expense_id, shop_id=shop_id).first()
            if not expense:
                return jsonify({'error': 'Expense not found'}), 404
            
            data = request.get_json()
            
            if 'item_name' in data:
                expense.item_name = data['item_name']
            if 'amount' in data:
                amount = float(data['amount'])
                if amount < 0:
                    return jsonify({'error': 'Amount cannot be negative'}), 400
                expense.amount = amount
            if 'date' in data:
                expense.date = datetime.strptime(data['date'], '%Y-%m-%d').date()
            if 'payment_method' in data:
                expense.payment_method = data['payment_method']
            if 'reference' in data:
                expense.reference = data['reference']
            if 'status' in data:
                expense.status = data['status']
            if 'notes' in data:
                expense.notes = data['notes']
            
            db.session.commit()
            return jsonify(expense.to_dict())
            
        except Exception as e:
            db.session.rollback()
            print(f"Error updating expense: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/expenses/<int:expense_id>', methods=['DELETE'])
    @login_required
    def delete_expense(expense_id):
        """Delete an expense for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            expense = Expense.query.filter_by(id=expense_id, shop_id=shop_id).first()
            if not expense:
                return jsonify({'error': 'Expense not found'}), 404
            
            db.session.delete(expense)
            db.session.commit()
            return jsonify({'message': 'Expense deleted successfully'})
            
        except Exception as e:
            db.session.rollback()
            print(f"Error deleting expense: {e}")
            return jsonify({'error': str(e)}), 500
    
    # ============ EXPENSE STATS ============
    
    @app.route('/api/expenses/stats', methods=['GET'])
    @login_required
    def get_expense_stats():
        """Get expense statistics for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            
            query = Expense.query.filter_by(shop_id=shop_id)
            
            if start_date:
                query = query.filter(Expense.date >= start_date)
            
            if end_date:
                query = query.filter(Expense.date <= end_date)
            
            expenses = query.all()
            
            total_expenses = sum(e.amount for e in expenses) if expenses else 0
            paid = sum(e.amount for e in expenses if e.status == 'Paid') if expenses else 0
            pending = sum(e.amount for e in expenses if e.status == 'Pending') if expenses else 0
            overdue = sum(e.amount for e in expenses if e.status == 'Overdue') if expenses else 0
            
            return jsonify({
                'total': float(total_expenses),
                'paid': float(paid),
                'pending': float(pending),
                'overdue': float(overdue),
                'count': len(expenses),
                'average': float(total_expenses / len(expenses)) if expenses and len(expenses) > 0 else 0
            })
            
        except Exception as e:
            print(f"Error fetching expense stats: {e}")
            return jsonify({'error': str(e)}), 500
    
    # ============ EXPORT ROUTES ============
    
    @app.route('/api/expenses/export/<string:format>', methods=['GET'])
    @login_required
    def export_expenses(format):
        """Export expenses data for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            status = request.args.get('status')
            payment_method = request.args.get('payment_method')
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            search = request.args.get('search')
            
            query = Expense.query.filter_by(shop_id=shop_id)
            
            if status:
                query = query.filter_by(status=status)
            
            if payment_method:
                query = query.filter_by(payment_method=payment_method)
            
            if start_date:
                query = query.filter(Expense.date >= start_date)
            
            if end_date:
                query = query.filter(Expense.date <= end_date)
            
            if search:
                query = query.filter(Expense.item_name.contains(search))
            
            expenses = query.order_by(Expense.date.desc()).all()
            
            if not expenses:
                return jsonify({
                    'error': 'No expenses data to export. Please create some expenses first.',
                    'message': 'You need to have expense records before exporting.'
                }), 404
            
            if format == 'excel':
                return export_excel(expenses)
            elif format == 'pdf':
                return export_pdf(expenses)
            elif format == 'word':
                return export_word(expenses)
            else:
                return jsonify({'error': 'Invalid format. Use pdf, excel, or word'}), 400
                
        except Exception as e:
            print(f"Error exporting expenses: {e}")
            return jsonify({'error': f'Failed to export expenses: {str(e)}'}), 500
    
    # ============ EXPORT HELPER FUNCTIONS ============
    
    def export_excel(expenses):
        """Export expenses to Excel"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Expenses Report"
        
        # Define styles
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="8B0000", end_color="8B0000", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        cell_alignment = Alignment(horizontal="left", vertical="center")
        number_alignment = Alignment(horizontal="right", vertical="center")
        center_alignment = Alignment(horizontal="center", vertical="center")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title row
        title_cell = ws.cell(row=1, column=1, value="EXPENSES REPORT")
        title_cell.font = Font(bold=True, size=16)
        ws.merge_cells('A1:I1')
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Date row
        date_cell = ws.cell(row=2, column=1, value=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        ws.merge_cells('A2:I2')
        date_cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Summary row
        total_expenses = sum(e.amount for e in expenses)
        paid = sum(e.amount for e in expenses if e.status == 'Paid')
        pending = sum(e.amount for e in expenses if e.status == 'Pending')
        overdue = sum(e.amount for e in expenses if e.status == 'Overdue')
        
        summary_row = 3
        ws.cell(row=summary_row, column=1, value=f"Total Expenses: {len(expenses)}")
        ws.cell(row=summary_row, column=3, value=f"Total Amount: KES {total_expenses:,.2f}")
        ws.cell(row=summary_row, column=5, value=f"Paid: KES {paid:,.2f}")
        ws.cell(row=summary_row, column=7, value=f"Pending: KES {pending:,.2f}")
        ws.cell(row=summary_row, column=9, value=f"Overdue: KES {overdue:,.2f}")
        
        # Headers (row 5)
        headers = [
            "ID", "Item Name", "Amount (KES)", "Date", 
            "Payment Method", "Reference", "Status", "Notes", "Created At"
        ]
        
        header_row = 5
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Data rows
        for row_idx, expense in enumerate(expenses, header_row + 1):
            row_data = [
                expense.id,
                expense.item_name,
                expense.amount,
                expense.date.strftime('%Y-%m-%d') if expense.date else 'N/A',
                expense.payment_method,
                expense.reference or '-',
                expense.status,
                expense.notes or '-',
                expense.created_at.strftime('%Y-%m-%d %H:%M') if expense.created_at else 'N/A'
            ]
            
            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col, value=value)
                cell.border = border
                if isinstance(value, (int, float)):
                    cell.alignment = number_alignment
                    if col == 3:  # Amount column
                        cell.number_format = '#,##0.00'
                elif col in [4, 9]:  # Date columns
                    cell.alignment = center_alignment
                else:
                    cell.alignment = cell_alignment
        
        # Auto-size columns
        for col in ws.columns:
            max_length = 0
            column = col[0].column
            for cell in col:
                try:
                    if isinstance(cell, openpyxl.cell.cell.MergedCell):
                        continue
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max(max_length + 2, 12), 30)
            ws.column_dimensions[get_column_letter(column)].width = adjusted_width
        
        # Save to BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'expenses_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    
    def export_pdf(expenses):
        """Export expenses to PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=landscape(letter),
            rightMargin=20,
            leftMargin=20,
            topMargin=20,
            bottomMargin=20
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=1,
            spaceAfter=10,
            textColor=colors.HexColor('#8B0000')
        )
        
        elements = []
        
        # Title
        title = Paragraph("Expenses Report", title_style)
        elements.append(title)
        
        # Date
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=1,
            spaceAfter=15
        )
        date_paragraph = Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", date_style)
        elements.append(date_paragraph)
        
        # Summary stats
        total_expenses = sum(e.amount for e in expenses)
        paid = sum(e.amount for e in expenses if e.status == 'Paid')
        pending = sum(e.amount for e in expenses if e.status == 'Pending')
        overdue = sum(e.amount for e in expenses if e.status == 'Overdue')
        
        summary_style = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=0,
            spaceAfter=15
        )
        summary_text = f"""
        <b>Summary:</b><br/>
        Total Expenses: {len(expenses)} | Total Amount: KES {total_expenses:,.2f} | 
        Paid: KES {paid:,.2f} | Pending: KES {pending:,.2f} | Overdue: KES {overdue:,.2f}
        """
        summary = Paragraph(summary_text, summary_style)
        elements.append(summary)
        
        # Prepare data for table
        table_data = [
            ['ID', 'Item Name', 'Amount', 'Date', 'Payment', 'Reference', 'Status']
        ]
        
        for expense in expenses:
            table_data.append([
                str(expense.id),
                expense.item_name[:25] + '...' if len(expense.item_name) > 25 else expense.item_name,
                f"KES {expense.amount:,.2f}",
                expense.date.strftime('%Y-%m-%d') if expense.date else 'N/A',
                expense.payment_method,
                expense.reference or '-',
                expense.status
            ])
        
        # Create table
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            # Header style
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B0000')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            # Alternating row colors
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('BACKGROUND', (0, 2), (-1, -1), colors.HexColor('#f8f9fa')),
            # Grid
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            # Alignment
            ('ALIGN', (2, 1), (2, -1), 'RIGHT'),
            ('ALIGN', (3, 1), (3, -1), 'CENTER'),
            ('ALIGN', (6, 1), (6, -1), 'CENTER'),
        ]))
        
        elements.append(table)
        elements.append(Spacer(1, 20))
        
        # Footer
        footer_style = ParagraphStyle(
            'FooterStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=1,
            textColor=colors.HexColor('#666666')
        )
        footer = Paragraph("Generated by POS System", footer_style)
        elements.append(footer)
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'expenses_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        )
    
    def export_word(expenses):
        """Export expenses to Word document"""
        doc = docx.Document()
        
        # Title
        title = doc.add_heading('Expenses Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(139, 0, 0)
        
        # Date
        date_paragraph = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Summary
        total_expenses = sum(e.amount for e in expenses)
        paid = sum(e.amount for e in expenses if e.status == 'Paid')
        pending = sum(e.amount for e in expenses if e.status == 'Pending')
        overdue = sum(e.amount for e in expenses if e.status == 'Overdue')
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Summary: ').bold = True
        summary_para.add_run(
            f'Total Expenses: {len(expenses)} | Total Amount: KES {total_expenses:,.2f} | '
            f'Paid: KES {paid:,.2f} | Pending: KES {pending:,.2f} | Overdue: KES {overdue:,.2f}'
        )
        
        doc.add_paragraph()
        
        # Add table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['ID', 'Item Name', 'Amount', 'Date', 'Payment', 'Reference', 'Status']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set header background
            shd_element = OxmlElement('w:shd')
            shd_element.set(qn('w:val'), 'clear')
            shd_element.set(qn('w:color'), 'auto')
            shd_element.set(qn('w:fill'), '8B0000')
            cell._element.tcPr.append(shd_element)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add data
        for expense in expenses:
            row_cells = table.add_row().cells
            row_cells[0].text = str(expense.id)
            row_cells[1].text = expense.item_name
            row_cells[2].text = f"KES {expense.amount:,.2f}"
            row_cells[3].text = expense.date.strftime('%Y-%m-%d') if expense.date else 'N/A'
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[4].text = expense.payment_method
            row_cells[5].text = expense.reference or '-'
            row_cells[6].text = expense.status
            row_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Auto-fit columns
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(1.0)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Generated by POS System').font.size = Pt(8)
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'expenses_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )