# routes/sales.py - COMPLETE IMPROVED FILE WITH SHOP NAME IN REPORTS (Tax Removed)
import uuid
from flask import request, jsonify, current_app, send_file
from flask_login import login_required, current_user
from extensions import db
from models.sales import Sale, SaleItem, SalePayment, Return, ReturnItem
from models.inventory import Product
from models.customer import Customer
from models.shop import Shop
from datetime import datetime
import io
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn, OxmlElement

def get_current_shop_id():
    """Get the current shop ID from the logged-in user"""
    if hasattr(current_user, 'id'):
        return current_user.id
    return None

def get_shop_name():
    """Get the current shop name from the logged-in user"""
    if hasattr(current_user, 'id'):
        shop = Shop.query.get(current_user.id)
        if shop:
            return shop.name
    return None

def init_sales_routes(app):
    
    # ============ SALES ROUTES ============
    
    @app.route('/api/sales', methods=['GET'])
    @login_required
    def get_sales():
        """Get all sales for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            status = request.args.get('status')
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            customer = request.args.get('customer')
            
            query = Sale.query.filter_by(shop_id=shop_id)
            
            if status:
                query = query.filter_by(status=status)
            
            if start_date:
                query = query.filter(Sale.created_at >= start_date)
            
            if end_date:
                end_datetime = datetime.strptime(end_date, '%Y-%m-%d')
                end_datetime = end_datetime.replace(hour=23, minute=59, second=59)
                query = query.filter(Sale.created_at <= end_datetime)
            
            if customer:
                query = query.filter(Sale.customer_name.contains(customer))
            
            sales = query.order_by(Sale.created_at.desc()).all()
            return jsonify([s.to_dict() for s in sales])
            
        except Exception as e:
            print(f"Error fetching sales: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/sales/<int:sale_id>', methods=['GET'])
    @login_required
    def get_sale(sale_id):
        """Get a specific sale by ID for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            sale = Sale.query.filter_by(id=sale_id, shop_id=shop_id).first()
            if not sale:
                return jsonify({'error': 'Sale not found'}), 404
            
            return jsonify(sale.to_dict())
            
        except Exception as e:
            print(f"Error fetching sale: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/sales', methods=['POST'])
    @login_required
    def create_sale():
        """Create a new sale for the current shop - Removed tax"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            data = request.get_json()
            
            # Validate required fields
            if not data or 'items' not in data or len(data['items']) == 0:
                return jsonify({'error': 'At least one item is required'}), 400
            
            # Validate items
            for item_data in data['items']:
                if 'product_id' not in item_data or 'quantity' not in item_data:
                    return jsonify({'error': 'Each item must have product_id and quantity'}), 400
                
                product = Product.query.filter_by(id=item_data['product_id'], shop_id=shop_id).first()
                if not product:
                    return jsonify({'error': f'Product {item_data["product_id"]} not found in this shop'}), 404
                
                if not product.is_active:
                    return jsonify({'error': f'Product {product.name} is inactive'}), 400
                
                if product.stock < item_data['quantity']:
                    return jsonify({'error': f'Insufficient stock for {product.name}. Available: {product.stock}'}), 400
            
            # Calculate totals - Removed tax
            subtotal = 0
            items = []
            
            for item_data in data['items']:
                product = Product.query.filter_by(id=item_data['product_id'], shop_id=shop_id).first()
                price = item_data.get('price', product.price)
                quantity = item_data['quantity']
                item_total = price * quantity
                subtotal += item_total
                
                items.append({
                    'product': product,
                    'price': price,
                    'quantity': quantity,
                    'total': item_total
                })
            
            discount = data.get('discount', 0)
            total = subtotal - discount  # Removed tax
            
            customer_id = data.get('customer_id')
            
            customer = None
            if customer_id:
                customer = Customer.query.filter_by(id=customer_id, shop_id=shop_id).first()
                if not customer:
                    return jsonify({'error': 'Customer not found in this shop'}), 404
            
            # Create sale with shop_id - Removed tax field
            sale = Sale(
                shop_id=shop_id,
                customer_id=customer_id if customer else None,
                customer_name=data.get('customer_name', 'Walk-in Customer'),
                customer_email=data.get('customer_email', ''),
                customer_phone=data.get('customer_phone', ''),
                payment_method=data.get('payment_method', 'Cash'),
                subtotal=subtotal,
                discount=discount,
                total=total,
                notes=data.get('notes', ''),
                status='Completed',
                created_by=current_user.id if current_user.is_authenticated else None
            )
            
            # Generate sale number
            sale.sale_number = sale.generate_sale_number()
            
            db.session.add(sale)
            db.session.flush()
            
            # Create sale items and update stock
            for item_data in items:
                product = item_data['product']
                
                sale_item = SaleItem(
                    shop_id=shop_id,
                    sale_id=sale.id,
                    product_id=product.id,
                    product_name=product.name,
                    product_sku=product.sku,
                    quantity=item_data['quantity'],
                    price=item_data['price'],
                    cost=product.cost,
                    total=item_data['total']
                )
                db.session.add(sale_item)
                
                product.stock -= item_data['quantity']
            
            # Create payment record with shop_id
            payment = SalePayment(
                shop_id=shop_id,
                sale_id=sale.id,
                payment_method=sale.payment_method,
                amount=total,
                status='Completed'
            )
            db.session.add(payment)
            
            # Update customer stats if customer is associated
            if customer:
                customer_sales = Sale.query.filter_by(
                    customer_id=customer.id,
                    shop_id=shop_id,
                    status='Completed'
                ).all()
                
                total_spent = sum(s.total for s in customer_sales)
                total_orders = len(customer_sales)
                
                customer.total_spent = total_spent
                customer.total_orders = total_orders
                customer.update_tier()
                customer.last_activity = datetime.utcnow()
            
            db.session.commit()
            
            return jsonify(sale.to_dict()), 201
            
        except Exception as e:
            db.session.rollback()
            print(f"Error creating sale: {e}")
            return jsonify({'error': f'Failed to create sale: {str(e)}'}), 500
    
    @app.route('/api/sales/<int:sale_id>', methods=['PUT'])
    @login_required
    def update_sale(sale_id):
        """Update a sale for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            sale = Sale.query.filter_by(id=sale_id, shop_id=shop_id).first()
            if not sale:
                return jsonify({'error': 'Sale not found'}), 404
            
            data = request.get_json()
            
            if 'customer_name' in data:
                sale.customer_name = data['customer_name']
            if 'customer_email' in data:
                sale.customer_email = data['customer_email']
            if 'customer_phone' in data:
                sale.customer_phone = data['customer_phone']
            if 'payment_method' in data:
                sale.payment_method = data['payment_method']
            if 'status' in data:
                valid_statuses = ['Completed', 'Pending', 'Cancelled', 'Refunded']
                if data['status'] not in valid_statuses:
                    return jsonify({'error': f'Invalid status. Must be one of: {", ".join(valid_statuses)}'}), 400
                sale.status = data['status']
            if 'notes' in data:
                sale.notes = data['notes']
            
            db.session.commit()
            return jsonify(sale.to_dict())
            
        except Exception as e:
            db.session.rollback()
            print(f"Error updating sale: {e}")
            return jsonify({'error': f'Failed to update sale: {str(e)}'}), 500
    
    @app.route('/api/sales/<int:sale_id>', methods=['DELETE'])
    @login_required
    def delete_sale(sale_id):
        """Delete a sale for the current shop (soft delete - mark as cancelled)"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            sale = Sale.query.filter_by(id=sale_id, shop_id=shop_id).first()
            if not sale:
                return jsonify({'error': 'Sale not found'}), 404
            
            sale.status = 'Cancelled'
            db.session.commit()
            
            return jsonify({'message': 'Sale cancelled successfully'})
            
        except Exception as e:
            db.session.rollback()
            print(f"Error cancelling sale: {e}")
            return jsonify({'error': f'Failed to cancel sale: {str(e)}'}), 500
    
    # ============ EXPORT ROUTES ============
    
    @app.route('/api/sales/export/<string:format>', methods=['GET'])
    @login_required
    def export_sales(format):
        """Export sales data for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            shop_name = get_shop_name() or "Sales Report"
            
            status = request.args.get('status')
            start_date = request.args.get('start_date')
            end_date = request.args.get('end_date')
            
            query = Sale.query.filter_by(shop_id=shop_id)
            
            if status:
                query = query.filter_by(status=status)
            
            if start_date:
                query = query.filter(Sale.created_at >= start_date)
            
            if end_date:
                end_datetime = datetime.strptime(end_date, '%Y-%m-%d')
                end_datetime = end_datetime.replace(hour=23, minute=59, second=59)
                query = query.filter(Sale.created_at <= end_datetime)
            
            sales = query.order_by(Sale.created_at.desc()).all()
            
            if not sales:
                return jsonify({
                    'error': 'No sales data to export. Please create some sales first.',
                    'message': 'You need to have sales records before exporting.'
                }), 404
            
            if format == 'excel':
                return export_excel(sales, shop_name)
            elif format == 'pdf':
                return export_pdf(sales, shop_name)
            elif format == 'word':
                return export_word(sales, shop_name)
            else:
                return jsonify({'error': 'Invalid format. Use pdf, excel, or word'}), 400
                
        except Exception as e:
            print(f"Error exporting sales: {e}")
            return jsonify({'error': f'Failed to export sales: {str(e)}'}), 500
    
    # ============ EXPORT HELPER FUNCTIONS - Removed tax from exports ============
    
    def export_excel(sales, shop_name):
        """Export sales to Excel with shop name in title - Removed tax"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sales Report"
        
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="1a5276", end_color="1a5276", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        cell_alignment = Alignment(horizontal="left", vertical="center")
        number_alignment = Alignment(horizontal="right", vertical="center")
        center_alignment = Alignment(horizontal="center", vertical="center")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title row with shop name
        title_cell = ws.cell(row=1, column=1, value=f"SALES REPORT - {shop_name.upper()}")
        title_cell.font = Font(bold=True, size=16)
        ws.merge_cells('A1:M1')
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Date row
        date_cell = ws.cell(row=2, column=1, value=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        ws.merge_cells('A2:M2')
        date_cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Summary row
        total_revenue = sum(s.total for s in sales)
        completed = len([s for s in sales if s.status == 'Completed'])
        pending = len([s for s in sales if s.status == 'Pending'])
        refunded = len([s for s in sales if s.status == 'Refunded'])
        total_items = sum(len(s.items) for s in sales)
        total_qty = sum(sum(item.quantity for item in s.items) for s in sales)
        
        summary_row = 3
        ws.cell(row=summary_row, column=1, value=f"Total Sales: {len(sales)}")
        ws.cell(row=summary_row, column=3, value=f"Total Revenue: KES {total_revenue:,.2f}")
        ws.cell(row=summary_row, column=5, value=f"Total Items: {total_items}")
        ws.cell(row=summary_row, column=7, value=f"Total Qty: {total_qty}")
        ws.cell(row=summary_row, column=9, value=f"Completed: {completed}")
        ws.cell(row=summary_row, column=11, value=f"Pending: {pending}")
        ws.cell(row=summary_row, column=13, value=f"Refunded: {refunded}")
        
        # Headers (row 5) - Removed tax column
        headers = [
            "Sale ID", "Customer", "Email", "Phone", "Date", 
            "Item Name", "SKU", "Item Qty", "Item Price", "Item Total",
            "Payment Method", "Order Total", "Status"
        ]
        
        header_row = 5
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Data rows
        row_idx = header_row + 1
        for sale in sales:
            if sale.items and len(sale.items) > 0:
                for item in sale.items:
                    row_data = [
                        sale.sale_number or f"S-{sale.id}",
                        sale.customer_name,
                        sale.customer_email or 'N/A',
                        sale.customer_phone or 'N/A',
                        sale.created_at.strftime('%Y-%m-%d %H:%M') if sale.created_at else 'N/A',
                        item.product_name or item.name,
                        item.product_sku or item.sku,
                        item.quantity,
                        item.price,
                        item.price * item.quantity,
                        sale.payment_method,
                        sale.total,
                        sale.status
                    ]
                    
                    for col, value in enumerate(row_data, 1):
                        cell = ws.cell(row=row_idx, column=col, value=value)
                        cell.border = border
                        if isinstance(value, (int, float)):
                            cell.alignment = number_alignment
                            if col in [9, 10, 12]:
                                cell.number_format = '#,##0.00'
                        elif col in [5, 8]:
                            cell.alignment = center_alignment
                        else:
                            cell.alignment = cell_alignment
                    
                    row_idx += 1
            else:
                row_data = [
                    sale.sale_number or f"S-{sale.id}",
                    sale.customer_name,
                    sale.customer_email or 'N/A',
                    sale.customer_phone or 'N/A',
                    sale.created_at.strftime('%Y-%m-%d %H:%M') if sale.created_at else 'N/A',
                    'No items',
                    'N/A',
                    0,
                    0,
                    0,
                    sale.payment_method,
                    sale.total,
                    sale.status
                ]
                
                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col, value=value)
                    cell.border = border
                    if isinstance(value, (int, float)):
                        cell.alignment = number_alignment
                    else:
                        cell.alignment = cell_alignment
                
                row_idx += 1
        
        # Auto-size columns
        for col in ws.columns:
            max_length = 0
            column = col[0].column
            for cell in col:
                try:
                    if isinstance(cell, openpyxl.cell.cell.MergedCell):
                        continue
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max(max_length + 2, 12), 30)
            ws.column_dimensions[get_column_letter(column)].width = adjusted_width
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'sales_report_{shop_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    
    def export_pdf(sales, shop_name):
        """Export sales to PDF with shop name in title - Removed tax"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=landscape(letter),
            rightMargin=20,
            leftMargin=20,
            topMargin=20,
            bottomMargin=20
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=1,
            spaceAfter=10,
            textColor=colors.HexColor('#1a5276')
        )
        
        elements = []
        
        title = Paragraph(f"{shop_name} - Sales Report", title_style)
        elements.append(title)
        
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=1,
            spaceAfter=15
        )
        date_paragraph = Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", date_style)
        elements.append(date_paragraph)
        
        total_revenue = sum(s.total for s in sales)
        completed = len([s for s in sales if s.status == 'Completed'])
        pending = len([s for s in sales if s.status == 'Pending'])
        refunded = len([s for s in sales if s.status == 'Refunded'])
        total_items = sum(len(s.items) for s in sales)
        total_qty = sum(sum(item.quantity for item in s.items) for s in sales)
        
        summary_style = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Normal'],
            fontSize=10,
            alignment=0,
            spaceAfter=15
        )
        summary_text = f"""
        <b>Summary:</b><br/>
        Total Sales: {len(sales)} | Total Revenue: KES {total_revenue:,.2f} | 
        Total Items: {total_items} | Total Qty: {total_qty} |
        Completed: {completed} | Pending: {pending} | Refunded: {refunded}
        """
        summary = Paragraph(summary_text, summary_style)
        elements.append(summary)
        
        # Removed tax column
        table_data = [
            ['Sale ID', 'Customer', 'Date', 'Item Name', 'SKU', 'Qty', 'Price', 'Item Total', 'Status']
        ]
        
        for sale in sales:
            if sale.items and len(sale.items) > 0:
                for item in sale.items:
                    table_data.append([
                        sale.sale_number or f"S-{sale.id}",
                        sale.customer_name[:20] + '...' if len(sale.customer_name) > 20 else sale.customer_name,
                        sale.created_at.strftime('%Y-%m-%d') if sale.created_at else 'N/A',
                        item.product_name or item.name,
                        item.product_sku or item.sku,
                        str(item.quantity),
                        f"KES {item.price:,.2f}",
                        f"KES {item.price * item.quantity:,.2f}",
                        sale.status
                    ])
            else:
                table_data.append([
                    sale.sale_number or f"S-{sale.id}",
                    sale.customer_name[:20] + '...' if len(sale.customer_name) > 20 else sale.customer_name,
                    sale.created_at.strftime('%Y-%m-%d') if sale.created_at else 'N/A',
                    'No items',
                    'N/A',
                    '0',
                    'KES 0.00',
                    'KES 0.00',
                    sale.status
                ])
        
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5276')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('BACKGROUND', (0, 2), (-1, -1), colors.HexColor('#f8f9fa')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cccccc')),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('ALIGN', (5, 1), (5, -1), 'CENTER'),
            ('ALIGN', (6, 1), (7, -1), 'RIGHT'),
            ('ALIGN', (8, 1), (8, -1), 'CENTER'),
        ]))
        
        elements.append(table)
        elements.append(Spacer(1, 20))
        
        footer_style = ParagraphStyle(
            'FooterStyle',
            parent=styles['Normal'],
            fontSize=8,
            alignment=1,
            textColor=colors.HexColor('#666666')
        )
        footer = Paragraph("Generated by POS System", footer_style)
        elements.append(footer)
        
        doc.build(elements)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'sales_report_{shop_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        )
    
    def export_word(sales, shop_name):
        """Export sales to Word document with shop name in title - Removed tax"""
        doc = docx.Document()
        
        title = doc.add_heading(f'{shop_name} - Sales Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        
        date_paragraph = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        total_revenue = sum(s.total for s in sales)
        completed = len([s for s in sales if s.status == 'Completed'])
        pending = len([s for s in sales if s.status == 'Pending'])
        refunded = len([s for s in sales if s.status == 'Refunded'])
        total_items = sum(len(s.items) for s in sales)
        total_qty = sum(sum(item.quantity for item in s.items) for s in sales)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Summary: ').bold = True
        summary_para.add_run(
            f'Total Sales: {len(sales)} | Total Revenue: KES {total_revenue:,.2f} | '
            f'Total Items: {total_items} | Total Qty: {total_qty} | '
            f'Completed: {completed} | Pending: {pending} | Refunded: {refunded}'
        )
        
        doc.add_paragraph()
        
        # Removed tax column (8 columns instead of 9)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        headers = ['Sale ID', 'Customer', 'Date', 'Item Name', 'SKU', 'Qty', 'Price', 'Item Total']
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shd_element = OxmlElement('w:shd')
            shd_element.set(qn('w:val'), 'clear')
            shd_element.set(qn('w:color'), 'auto')
            shd_element.set(qn('w:fill'), '1a5276')
            cell._element.tcPr.append(shd_element)
        
        for sale in sales:
            if sale.items and len(sale.items) > 0:
                for item in sale.items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = sale.sale_number or f"S-{sale.id}"
                    row_cells[1].text = sale.customer_name
                    row_cells[2].text = sale.created_at.strftime('%Y-%m-%d') if sale.created_at else 'N/A'
                    row_cells[3].text = item.product_name or item.name
                    row_cells[4].text = item.product_sku or item.sku
                    row_cells[5].text = str(item.quantity)
                    row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[6].text = f"KES {item.price:,.2f}"
                    row_cells[7].text = f"KES {item.price * item.quantity:,.2f}"
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = sale.sale_number or f"S-{sale.id}"
                row_cells[1].text = sale.customer_name
                row_cells[2].text = sale.created_at.strftime('%Y-%m-%d') if sale.created_at else 'N/A'
                row_cells[3].text = 'No items'
                row_cells[4].text = 'N/A'
                row_cells[5].text = '0'
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[6].text = 'KES 0.00'
                row_cells[7].text = 'KES 0.00'
        
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(1.0)
        
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('Generated by POS System').font.size = Pt(8)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'sales_report_{shop_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
    
    # ============ RETURNS ROUTES ============
    
    @app.route('/api/returns', methods=['GET'])
    @login_required
    def get_returns():
        """Get all returns for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            status = request.args.get('status')
            
            query = Return.query.filter_by(shop_id=shop_id)
            if status:
                query = query.filter_by(status=status)
            
            returns = query.order_by(Return.created_at.desc()).all()
            return jsonify([r.to_dict() for r in returns])
            
        except Exception as e:
            print(f"Error fetching returns: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/returns', methods=['POST'])
    @login_required
    def create_return():
        """Create a new return for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            data = request.get_json()
            
            if not data or 'sale_id' not in data:
                return jsonify({'error': 'Sale ID is required'}), 400
            
            sale = Sale.query.filter_by(id=data['sale_id'], shop_id=shop_id).first()
            if not sale:
                return jsonify({'error': 'Sale not found in this shop'}), 404
            
            if 'items' not in data or len(data['items']) == 0:
                return jsonify({'error': 'At least one item is required for return'}), 400
            
            total_refund = 0
            return_items = []
            
            for item_data in data['items']:
                product = Product.query.filter_by(id=item_data['product_id'], shop_id=shop_id).first()
                if not product:
                    return jsonify({'error': f'Product {item_data["product_id"]} not found in this shop'}), 404
                
                quantity = item_data['quantity']
                refund_amount = item_data.get('refund_amount', product.price * quantity)
                total_refund += refund_amount
                
                return_items.append({
                    'product': product,
                    'quantity': quantity,
                    'price': product.price,
                    'refund_amount': refund_amount,
                    'reason': item_data.get('reason', ''),
                    'condition': item_data.get('condition', 'Good')
                })
            
            return_record = Return(
                shop_id=shop_id,
                return_number=Return.generate_return_number(Return),
                sale_id=sale.id,
                customer_name=data.get('customer_name', sale.customer_name),
                customer_email=data.get('customer_email', sale.customer_email),
                customer_phone=data.get('customer_phone', sale.customer_phone),
                reason=data.get('reason', ''),
                total=total_refund,
                notes=data.get('notes', ''),
                status='Pending',
                created_by=current_user.id if current_user.is_authenticated else None
            )
            
            db.session.add(return_record)
            db.session.flush()
            
            for item_data in return_items:
                product = item_data['product']
                
                return_item = ReturnItem(
                    shop_id=shop_id,
                    return_id=return_record.id,
                    product_id=product.id,
                    product_name=product.name,
                    product_sku=product.sku,
                    quantity=item_data['quantity'],
                    price=item_data['price'],
                    refund_amount=item_data['refund_amount'],
                    reason=item_data['reason'],
                    condition=item_data['condition']
                )
                db.session.add(return_item)
                
                product.stock += item_data['quantity']
            
            db.session.commit()
            
            return jsonify(return_record.to_dict()), 201
            
        except Exception as e:
            db.session.rollback()
            print(f"Error creating return: {e}")
            return jsonify({'error': f'Failed to create return: {str(e)}'}), 500
    
    @app.route('/api/returns/<int:return_id>', methods=['PUT'])
    @login_required
    def update_return(return_id):
        """Update a return status for the current shop"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            return_record = Return.query.filter_by(id=return_id, shop_id=shop_id).first()
            if not return_record:
                return jsonify({'error': 'Return not found'}), 404
            
            data = request.get_json()
            
            valid_statuses = ['Pending', 'Approved', 'Rejected', 'Completed']
            if 'status' in data:
                if data['status'] not in valid_statuses:
                    return jsonify({'error': f'Invalid status. Must be one of: {", ".join(valid_statuses)}'}), 400
                
                return_record.status = data['status']
            
            if 'notes' in data:
                return_record.notes = data['notes']
            
            db.session.commit()
            return jsonify(return_record.to_dict())
            
        except Exception as e:
            db.session.rollback()
            print(f"Error updating return: {e}")
            return jsonify({'error': f'Failed to update return: {str(e)}'}), 500
    
    # ============ DASHBOARD STATS ============
    
    @app.route('/api/sales/stats', methods=['GET'])
    @login_required
    def get_sales_stats():
        """Get sales statistics for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
            
            today_sales = Sale.query.filter(
                Sale.shop_id == shop_id,
                Sale.created_at >= today
            ).all()
            today_revenue = sum(s.total for s in today_sales)
            today_count = len(today_sales)
            
            month_start = today.replace(day=1)
            month_sales = Sale.query.filter(
                Sale.shop_id == shop_id,
                Sale.created_at >= month_start
            ).all()
            month_revenue = sum(s.total for s in month_sales)
            month_count = len(month_sales)
            
            all_sales = Sale.query.filter_by(shop_id=shop_id).all()
            total_revenue = sum(s.total for s in all_sales)
            total_count = len(all_sales)
            
            completed = Sale.query.filter_by(shop_id=shop_id, status='Completed').count()
            pending = Sale.query.filter_by(shop_id=shop_id, status='Pending').count()
            refunded = Sale.query.filter_by(shop_id=shop_id, status='Refunded').count()
            cancelled = Sale.query.filter_by(shop_id=shop_id, status='Cancelled').count()
            
            return jsonify({
                'today': {
                    'revenue': float(today_revenue),
                    'count': today_count
                },
                'month': {
                    'revenue': float(month_revenue),
                    'count': month_count
                },
                'total': {
                    'revenue': float(total_revenue),
                    'count': total_count
                },
                'status_breakdown': {
                    'completed': completed,
                    'pending': pending,
                    'refunded': refunded,
                    'cancelled': cancelled
                }
            })
            
        except Exception as e:
            print(f"Error fetching sales stats: {e}")
            return jsonify({'error': str(e)}), 500
    
    @app.route('/api/sales/recent', methods=['GET'])
    @login_required
    def get_recent_sales():
        """Get recent sales for the current shop only"""
        try:
            shop_id = get_current_shop_id()
            if not shop_id:
                return jsonify({'error': 'Shop not found'}), 401
            
            limit = request.args.get('limit', 10, type=int)
            sales = Sale.query.filter_by(shop_id=shop_id).order_by(Sale.created_at.desc()).limit(limit).all()
            return jsonify([s.to_dict() for s in sales])
            
        except Exception as e:
            print(f"Error fetching recent sales: {e}")
            return jsonify({'error': str(e)}), 500